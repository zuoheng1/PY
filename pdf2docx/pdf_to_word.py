#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF转Word工具
支持多种转换方式：pdf2docx、python-docx + PyPDF2、以及OCR识别
"""

import os
import sys
from pathlib import Path
import argparse
from typing import List, Optional
import logging
from datetime import datetime

# 设置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('pdf_conversion.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

class PDFToWordConverter:
    """PDF转Word转换器"""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # 支持的PDF文件扩展名
        self.supported_extensions = {'.pdf'}
        
        # 转换统计
        self.stats = {
            'total': 0,
            'success': 0,
            'failed': 0,
            'skipped': 0
        }
    
    def convert_with_pdf2docx(self, pdf_path: Path, output_path: Path) -> bool:
        """使用pdf2docx库进行转换（推荐方法）"""
        try:
            from pdf2docx import Converter
            
            logger.info(f"🔄 使用pdf2docx转换: {pdf_path.name}")
            
            cv = Converter(str(pdf_path))
            cv.convert(str(output_path), start=0, end=None)
            cv.close()
            
            logger.info(f"✅ pdf2docx转换成功: {output_path.name}")
            return True
            
        except ImportError:
            logger.warning("❌ pdf2docx库未安装，请运行: pip install pdf2docx")
            return False
        except Exception as e:
            logger.error(f"❌ pdf2docx转换失败 {pdf_path.name}: {e}")
            return False
    
    def convert_with_pypdf_docx(self, pdf_path: Path, output_path: Path) -> bool:
        """使用PyPDF2 + python-docx进行转换"""
        try:
            import PyPDF2
            from docx import Document
            
            logger.info(f"🔄 使用PyPDF2+docx转换: {pdf_path.name}")
            
            # 读取PDF
            with open(pdf_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                # 创建Word文档
                doc = Document()
                doc.add_heading(f'转换自: {pdf_path.name}', 0)
                
                # 提取每页文本
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            doc.add_heading(f'第 {page_num} 页', level=1)
                            doc.add_paragraph(text)
                        else:
                            doc.add_heading(f'第 {page_num} 页', level=1)
                            doc.add_paragraph('[此页面无法提取文本，可能包含图片或特殊格式]')
                    except Exception as e:
                        logger.warning(f"⚠️  第{page_num}页提取失败: {e}")
                        doc.add_heading(f'第 {page_num} 页', level=1)
                        doc.add_paragraph(f'[页面提取错误: {e}]')
                
                # 保存Word文档
                doc.save(str(output_path))
            
            logger.info(f"✅ PyPDF2+docx转换成功: {output_path.name}")
            return True
            
        except ImportError as e:
            logger.warning(f"❌ 缺少依赖库: {e}")
            return False
        except Exception as e:
            logger.error(f"❌ PyPDF2+docx转换失败 {pdf_path.name}: {e}")
            return False
    
    def convert_with_ocr(self, pdf_path: Path, output_path: Path) -> bool:
        """使用OCR进行转换（适用于扫描版PDF）"""
        try:
            import fitz  # PyMuPDF
            import pytesseract
            from PIL import Image
            from docx import Document
            import io
            
            logger.info(f"🔄 使用OCR转换: {pdf_path.name}")
            
            # 打开PDF
            pdf_document = fitz.open(str(pdf_path))
            doc = Document()
            doc.add_heading(f'OCR转换自: {pdf_path.name}', 0)
            
            for page_num in range(len(pdf_document)):
                try:
                    page = pdf_document.load_page(page_num)
                    
                    # 将页面转换为图片
                    mat = fitz.Matrix(2.0, 2.0)  # 提高分辨率
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    
                    # 使用PIL打开图片
                    image = Image.open(io.BytesIO(img_data))
                    
                    # OCR识别
                    text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                    
                    if text.strip():
                        doc.add_heading(f'第 {page_num + 1} 页', level=1)
                        doc.add_paragraph(text)
                    else:
                        doc.add_heading(f'第 {page_num + 1} 页', level=1)
                        doc.add_paragraph('[OCR未识别到文本内容]')
                        
                except Exception as e:
                    logger.warning(f"⚠️  第{page_num + 1}页OCR失败: {e}")
                    doc.add_heading(f'第 {page_num + 1} 页', level=1)
                    doc.add_paragraph(f'[OCR识别错误: {e}]')
            
            pdf_document.close()
            doc.save(str(output_path))
            
            logger.info(f"✅ OCR转换成功: {output_path.name}")
            return True
            
        except ImportError as e:
            logger.warning(f"❌ OCR转换缺少依赖: {e}")
            logger.info("请安装: pip install PyMuPDF pytesseract pillow")
            return False
        except Exception as e:
            logger.error(f"❌ OCR转换失败 {pdf_path.name}: {e}")
            return False
    
    def convert_single_file(self, pdf_path: Path, method: str = 'auto') -> bool:
        """转换单个PDF文件"""
        if not pdf_path.exists():
            logger.error(f"❌ 文件不存在: {pdf_path}")
            return False
        
        if pdf_path.suffix.lower() not in self.supported_extensions:
            logger.warning(f"⚠️  跳过非PDF文件: {pdf_path.name}")
            self.stats['skipped'] += 1
            return False
        
        # 生成输出文件名
        output_filename = pdf_path.stem + '.docx'
        output_path = self.output_dir / output_filename
        
        # 如果文件已存在，添加时间戳
        if output_path.exists():
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"{pdf_path.stem}_{timestamp}.docx"
            output_path = self.output_dir / output_filename
        
        self.stats['total'] += 1
        
        # 根据方法选择转换方式
        success = False
        
        if method == 'auto':
            # 自动选择最佳方法
            success = (self.convert_with_pdf2docx(pdf_path, output_path) or
                      self.convert_with_pypdf_docx(pdf_path, output_path) or
                      self.convert_with_ocr(pdf_path, output_path))
        elif method == 'pdf2docx':
            success = self.convert_with_pdf2docx(pdf_path, output_path)
        elif method == 'pypdf':
            success = self.convert_with_pypdf_docx(pdf_path, output_path)
        elif method == 'ocr':
            success = self.convert_with_ocr(pdf_path, output_path)
        else:
            logger.error(f"❌ 未知的转换方法: {method}")
            return False
        
        if success:
            self.stats['success'] += 1
            logger.info(f"🎉 转换完成: {pdf_path.name} -> {output_filename}")
        else:
            self.stats['failed'] += 1
            logger.error(f"❌ 转换失败: {pdf_path.name}")
        
        return success
    
    def convert_directory(self, input_dir: Path, method: str = 'auto', recursive: bool = False) -> None:
        """转换目录中的所有PDF文件"""
        if not input_dir.exists():
            logger.error(f"❌ 目录不存在: {input_dir}")
            return
        
        logger.info(f"🚀 开始批量转换: {input_dir}")
        logger.info(f"📁 输出目录: {self.output_dir}")
        logger.info(f"🔧 转换方法: {method}")
        logger.info(f"🔄 递归搜索: {'是' if recursive else '否'}")
        
        # 查找PDF文件
        pattern = '**/*.pdf' if recursive else '*.pdf'
        pdf_files = list(input_dir.glob(pattern))
        
        if not pdf_files:
            logger.warning(f"⚠️  未找到PDF文件: {input_dir}")
            return
        
        logger.info(f"📋 找到 {len(pdf_files)} 个PDF文件")
        
        # 转换每个文件
        for pdf_file in pdf_files:
            logger.info(f"\n{'='*60}")
            self.convert_single_file(pdf_file, method)
        
        # 打印统计信息
        self.print_stats()
    
    def print_stats(self) -> None:
        """打印转换统计信息"""
        logger.info(f"\n{'='*60}")
        logger.info("📊 转换统计:")
        logger.info(f"   总文件数: {self.stats['total']}")
        logger.info(f"   成功转换: {self.stats['success']}")
        logger.info(f"   转换失败: {self.stats['failed']}")
        logger.info(f"   跳过文件: {self.stats['skipped']}")
        
        if self.stats['total'] > 0:
            success_rate = (self.stats['success'] / self.stats['total']) * 100
            logger.info(f"   成功率: {success_rate:.1f}%")

def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description='PDF转Word工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  python pdf_to_word.py file.pdf                    # 转换单个文件
  python pdf_to_word.py /path/to/pdfs/              # 转换目录中的所有PDF
  python pdf_to_word.py /path/to/pdfs/ -r           # 递归转换子目录
  python pdf_to_word.py file.pdf -m ocr            # 使用OCR方法
  python pdf_to_word.py file.pdf -o /output/dir    # 指定输出目录

转换方法:
  auto     - 自动选择最佳方法（默认）
  pdf2docx - 使用pdf2docx库（推荐，保持格式）
  pypdf    - 使用PyPDF2+python-docx（纯文本）
  ocr      - 使用OCR识别（适用于扫描版PDF）
        """
    )
    
    parser.add_argument('input', help='输入PDF文件或目录路径')
    parser.add_argument('-o', '--output', default='output', help='输出目录（默认: output）')
    parser.add_argument('-m', '--method', choices=['auto', 'pdf2docx', 'pypdf', 'ocr'], 
                       default='auto', help='转换方法（默认: auto）')
    parser.add_argument('-r', '--recursive', action='store_true', help='递归处理子目录')
    parser.add_argument('-v', '--verbose', action='store_true', help='详细输出')
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # 创建转换器
    converter = PDFToWordConverter(args.output)
    
    input_path = Path(args.input)
    
    if input_path.is_file():
        # 转换单个文件
        converter.convert_single_file(input_path, args.method)
    elif input_path.is_dir():
        # 转换目录
        converter.convert_directory(input_path, args.method, args.recursive)
    else:
        logger.error(f"❌ 输入路径无效: {input_path}")
        sys.exit(1)

if __name__ == '__main__':
    main()